"""Artifact parsing — turn uploaded PDF/DOCX (or raw text) into plain text.

Scope: pure text extraction. No layout preservation, no image OCR. The four
artifact slots are small — a values document, a role spec, a team-structure
blurb, a sample comms excerpt. Paste-as-text is the primary path; uploads are
a convenience.
"""
from __future__ import annotations

import io
from pathlib import PurePath

from docx import Document
from pypdf import PdfReader


class UnsupportedArtifactType(ValueError):
    pass


def parse_upload(*, filename: str, data: bytes) -> str:
    """Dispatch on file extension."""
    ext = PurePath(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext in {".txt", ".md", ""}:
        # Plain text / markdown — decode and return.
        return data.decode("utf-8", errors="replace")
    raise UnsupportedArtifactType(
        f"Unsupported artifact file type: {ext or '(no extension)'}. "
        f"Accepted: .pdf, .docx, .txt, .md"
    )


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            # pypdf occasionally throws on pathological pages — skip them
            # rather than failing the whole upload.
            continue
    return "\n\n".join(c.strip() for c in chunks if c.strip())


def _parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Pull table cell text too — role specs often live in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
